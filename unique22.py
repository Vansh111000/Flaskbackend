import google.generativeai as genai
import json
import numpy as np
from docx import Document
from datetime import datetime
import os

# Configure the API key for Google Gemini

genai.configure(api_key="AIzaSyCckJqoHEUlJmOUDstkiTPG0p-0Xru9Xyo")  # Replace with your actual API key

# Use Gemini 1.5 Flash model
model = genai.GenerativeModel("gemini-1.5-flash")

# Function to get top 3 most relevant diseases based on symptoms
def get_relevant_diseases(symptoms_list):
    if not symptoms_list:
        return []

    prompt = f"""
    Based on the following symptoms: {', '.join(symptoms_list)},
    predict the **top 3 most likely diseases** that match these symptoms.

    For each disease, provide:
    - 🧠 Disease Name
    - 📊 Probability (High, Medium, Low)
    - 📝 Reason (explain which symptoms led to this prediction)
    - ⚠️ Precautions
    - 🛠️ Remedies

    Only list **3 diseases** in a structured JSON format as follows:
    [
        {{
            "disease": "Disease Name",
            "probability": "High/Medium/Low",
            "reason": "Brief explanation",
            "precautions": "Precautionary measures",
            "remedies": "Suggested remedies"
        }},
        ...
    ]
    """

    for _ in range(3):  # Retry logic in case of errors
        response = model.generate_content(prompt)
        try:
            json_text = response.text.strip().strip("```json").strip("```")
            diseases = json.loads(json_text)
            if isinstance(diseases, list) and len(diseases) == 3:
                return diseases
        except json.JSONDecodeError:
            continue  # Retry if parsing fails
    return []  # Return empty list if all retries fail

# Function to generate follow-up questions based on a symptom
def get_symptom_questions(symptom):
    prompt = f"""
    Based on the symptom '{symptom}', generate **7 relevant follow-up symptom-related questions**.
    - Each question must have 3 answer choices: ["No", "Mild", "Severe"].
    - The questions should help in diagnosing diseases accurately.
    - Format the response as a valid JSON list without any extra text.

    Example:
    [
        {{
            "question": "How severe is the pain?",
            "answers": ["No", "Mild", "Severe"]
        }},
        ...
    ]
    """

    for _ in range(3):  # Retry logic
        response = model.generate_content(prompt)
        try:
            json_text = response.text.strip().strip("```json").strip("```")
            questions = json.loads(json_text)
            if isinstance(questions, list) and len(questions) >= 5:
                return questions[:7]  # Ensure we get exactly 7 questions
        except json.JSONDecodeError:
            continue  # Retry if JSON parsing fails
    return []  # Return empty list if all retries fail

# Function to get user responses
def get_user_responses(questions):
    responses = []
    user_answers = []
    choices = {"A": 0, "B": 1, "C": 2}
    
    print("\n🔍 Please answer the following questions:")
    for i, q in enumerate(questions):
        print(f"{i+1}. {q['question']}")
        print(f"(A) {q['answers'][0]}  (B) {q['answers'][1]}  (C) {q['answers'][2]}")
        while True:
            user_input = input("Your choice (A/B/C): ").strip().upper()
            if user_input in choices:
                responses.append(choices[user_input])
                user_answers.append({"question": q["question"], "answer": q["answers"][choices[user_input]]})
                break
            else:
                print("⚠️ Invalid choice! Please enter A, B, or C.")
    
    return np.array(responses).reshape(1, -1), user_answers

# Function to generate the medical report
def generate_report(user_answers, symptoms_list, predicted_diseases):
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    file_name = f"Medical_Report_{timestamp}.docx"
    doc = Document()
    doc.add_heading("Medical Symptom Assessment Report", level=1)
    
    doc.add_heading("Patient's Symptom Responses:", level=2)
    for entry in user_answers:
        doc.add_paragraph(f"- {entry['question']}: {entry['answer']}")
    
    if symptoms_list:
        doc.add_heading("Additional Symptoms:", level=2)
        for symptom in symptoms_list:
            doc.add_paragraph(f"- {symptom}")
    
    doc.add_heading("Predicted Conditions:", level=2)
    for disease in predicted_diseases:
        doc.add_paragraph(f"- 🧠 {disease['disease']} - **{disease['probability']} Probability**")
        doc.add_paragraph(f"   📝 {disease['reason']}")
        doc.add_paragraph(f"   ⚠️ Precautions: {disease['precautions']}")
        doc.add_paragraph(f"   🛠️ Remedies: {disease['remedies']}")
    
    doc.add_paragraph("\n⚠️ This is not a final diagnosis. Please consult a doctor.")
    doc.save(file_name)
    print(f"📄 Report generated: {file_name}")

# Main disease prediction function
def predict_disease():
    # Reset previous data on each run
    symptoms_list = []

    additional_symptoms = input("\nPlease describe any additional symptoms you're experiencing (or press Enter to skip): ")
    symptoms_list = [s.strip() for s in additional_symptoms.split(',') if s.strip()]
    
    all_questions = []
    seen_questions = set()

    # Get follow-up questions based on symptoms
    for symptom in symptoms_list:
        questions = get_symptom_questions(symptom)
        for q in questions:
            if q["question"] not in seen_questions:
                all_questions.append(q)
                seen_questions.add(q["question"])

    # Get user responses
    user_data, user_answers = get_user_responses(all_questions)

    # Get top 3 disease predictions
    predicted_diseases = get_relevant_diseases(symptoms_list)

    # Display predicted diseases
    if predicted_diseases:
        print("\n🩺 **Predicted Diseases:**")
        for i, disease in enumerate(predicted_diseases):
            print(f"{i+1}. 🧠 {disease['disease']} - **{disease['probability']} Probability**")
            print(f"   📝 {disease['reason']}")
            print(f"   ⚠️ Precautions: {disease['precautions']}")
            print(f"   🛠️ Remedies: {disease['remedies']}\n")
    else:
        print("⚠️ No specific diseases could be determined. Please consult a doctor for further evaluation.")

    # Generate report
    generate_report(user_answers, symptoms_list, predicted_diseases)

# Run the prediction function
predict_disease()
