from flask import Flask, request, jsonify, send_from_directory
import google.generativeai as genai
import json
import os
import numpy as np
from docx import Document
from datetime import datetime
from flask_cors import CORS

# Flask app setup
app = Flask(__name__)
CORS(app)

# Configure Gemini API
genai.configure(api_key="AIzaSyCIhLsGzLC2KhzMOGmAvZqt4pj5FXSilY0")  # Replace with your key
model = genai.GenerativeModel("gemini-1.5-flash")

REPORTS_DIR = "reports"
os.makedirs(REPORTS_DIR, exist_ok=True)

# Endpoint 1: Get follow-up questions for symptoms
@app.route("/questions", methods=["POST"])
def get_questions():
    data = request.get_json()
    symptoms = data.get("symptoms", [])
    all_questions = []
    seen = set()

    for symptom in symptoms:
        prompt = f"""
        Based on the symptom '{symptom}', generate 7 follow-up questions with options ["No", "Mild", "Severe"].
        Format the response as JSON:
        [
            {{
                "question": "text",
                "answers": ["No", "Mild", "Severe"]
            }},
            ...
        ]
        """
        for _ in range(2):
            try:
                response = model.generate_content(prompt)
                json_text = response.text.strip().strip("```json").strip("```")
                questions = json.loads(json_text)
                for q in questions:
                    if q["question"] not in seen:
                        all_questions.append(q)
                        seen.add(q["question"])
                break
            except Exception:
                continue

    return jsonify(all_questions[:7])

# Endpoint 2: Predict diseases based on symptoms
@app.route("/predict", methods=["POST"])
def predict_disease():
    data = request.get_json()
    print("Symptoms data::: ",data)
    symptoms = data.get("symptoms", [])
    prompt = f"""
    Given these symptoms: {', '.join(symptoms)},
    return a JSON list of the top 3 diseases like:
    [
        {{
            "disease": "Name",
            "probability": "High/Medium/Low",
            "reason": "Why it's predicted",
            "precautions": "Preventive measures",
            "remedies": "Suggestions"
        }}
    ]
    """

    print("Symptoms data::: after prompt ",data)

    for _ in range(2):
        try:
            print("Symptoms data::: In try method",data)
            response = model.generate_content(prompt)
            json_text = response.text.strip().strip("```json").strip("```")
            print("jsonText loaded successfuly",json_text)
            diseases = json.loads(json_text)
            print("diseases loaded successfully",len(diseases))
            if isinstance(diseases, list) and len(diseases) == 3:
                return jsonify(diseases)
        except Exception:
            continue

    return jsonify([]), 400

# Endpoint 3: Generate report and return download link
@app.route("/report", methods=["POST"])
def generate_report():
    data = request.get_json()
    answers = data.get("answers", [])
    symptoms = data.get("symptoms", [])
    predictions = data.get("predictions", [])

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"Medical_Report_{timestamp}.docx"
    filepath = os.path.join(REPORTS_DIR, filename)

    doc = Document()
    doc.add_heading("Medical Symptom Assessment Report", level=1)

    doc.add_heading("Symptom Responses", level=2)
    for a in answers:
        doc.add_paragraph(f"- {a['question']}: {a['answer']}")

    if symptoms:
        doc.add_heading("Additional Symptoms", level=2)
        for s in symptoms:
            doc.add_paragraph(f"- {s}")

    doc.add_heading("Predicted Diseases", level=2)
    for d in predictions:
        doc.add_paragraph(f"- 🧠 {d['disease']} - {d['probability']} Probability")
        doc.add_paragraph(f"   📝 {d['reason']}")
        doc.add_paragraph(f"   ⚠️ Precautions: {d['precautions']}")
        doc.add_paragraph(f"   🛠️ Remedies: {d['remedies']}")

    doc.add_paragraph("\n⚠️ This is not a diagnosis. Please consult a doctor.")
    doc.save(filepath)

    return jsonify({"file": f"/download/{filename}"}), 200

# Endpoint to download the generated report
@app.route("/download/<filename>")
def download_report(filename):
    return send_from_directory(REPORTS_DIR, filename, as_attachment=True)

# Run the Flask app
if __name__ == "__main__":
    app.run(debug=True)
